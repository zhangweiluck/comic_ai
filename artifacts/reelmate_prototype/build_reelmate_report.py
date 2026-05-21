from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/zhangwei/JSWorkSpace/learning_code/comic_ai")
OUT_DIR = ROOT / "artifacts" / "reelmate_prototype"
IMG_DIR = OUT_DIR / "screenshots"
DOCX_PATH = OUT_DIR / "ReelMate_AI_Comic_Drama_Prototype_Report.docx"
VERIFIED_SCREENSHOTS = {
    "00_home_logged_in.png",
    "01_script_manage_empty.png",
    "12_script_ai_original_settings_modal.png",
    "13_script_ai_original_filled_disabled.png",
    "14_script_episode_count_dropdown.png",
    "03_project_list_with_existing_project.png",
    "15_project_create_modal.png",
    "16_project_create_validation_toast.png",
    "17_project_card_more_menu.png",
    "18_project_rename_modal.png",
    "05_project_overview_assets_and_episodes.png",
    "19_project_script_upload_modal_from_extract.png",
    "20_project_script_library_empty_in_upload_modal.png",
    "21_project_storyboard_upload_tab.png",
    "06_episode_storyboard_editor_guided_tour.png",
    "22_episode_video_model_dropdown.png",
    "23_episode_generate_validation_toast.png",
    "24_episode_add_storyboard_result.png",
    "07_user_asset_library_empty.png",
    "08_team_member_management_gate.png",
    "09_team_dashboard_gate.png",
    "10_team_knowledge_base_gate.png",
    "25_team_create_member_pricing_gate.png",
    "26_team_member_rules_modal.png",
}


COLORS = {
    "navy": "1F2937",
    "purple": "7C3AED",
    "lavender": "F3EEFF",
    "gray": "F3F4F6",
    "line": "D1D5DB",
    "muted": "6B7280",
    "white": "FFFFFF",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_east_asia_font(run, "Microsoft YaHei")


def set_cell_border(cell, color="D1D5DB", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_east_asia_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def style_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["navy"])

    for name, size, color in [
        ("Title", 26, COLORS["navy"]),
        ("Heading 1", 17, COLORS["purple"]),
        ("Heading 2", 13, COLORS["navy"]),
        ("Heading 3", 11, COLORS["navy"]),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_meta_table(doc):
    rows = [
        ("采集对象", "ReelMate / 万兴剧厂，AI驱动的一站式精品漫剧创作平台"),
        ("入口 URL", "https://www.reelmate.cn/home"),
        ("采集日期", "2026-05-19"),
        ("采集环境", "已登录 Chrome 账号，桌面视口；账号积分为 0，团队专业版权益/席位受限"),
        ("报告版本", "v1.1 强交互触发态补充稿"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, color=COLORS["white"])
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], COLORS["purple"])
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_note(doc, text, fill="F9FAFB"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "E5E7EB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    set_east_asia_font(run, "Microsoft YaHei")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_east_asia_font(run, "Microsoft YaHei")


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_east_asia_font(run, "Microsoft YaHei")


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=COLORS["white"])
        set_cell_shading(cell, COLORS["purple"])
        set_cell_border(cell)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            set_cell_border(cells[i])
    doc.add_paragraph()
    return table


def add_screenshot(doc, filename, title, details):
    path = IMG_DIR / filename
    doc.add_heading(title, level=3)
    if filename in VERIFIED_SCREENSHOTS and path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.75))
    else:
        add_note(doc, f"截图待补采：{filename}。本页旅途与控件来自已登录 Chrome 的实时页面观察；当前账号/自动化环境未能稳定打开或保存该状态，因此未嵌入未验证图片。", "FEF2F2")
    headers = ["说明项", "内容"]
    rows = [
        ("页面目的", details["purpose"]),
        ("触发入口", details["entry"]),
        ("主要控件", details["controls"]),
        ("下一步", details["next"]),
        ("复刻要点", details["prototype"]),
    ]
    add_simple_table(doc, headers, rows)


def add_cover(doc):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ReelMate AI漫剧产品原型采集报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    set_east_asia_font(run, "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("剧本、项目、资产库、团队四条用户旅途复刻依据")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    set_east_asia_font(run, "Microsoft YaHei")
    doc.add_paragraph()
    add_meta_table(doc)
    doc.add_paragraph()
    add_note(
        doc,
        "本报告基于已登录账号的实际可见页面与公开前端路由信息整理。当前账号积分为 0，且团队资产库/团队席位存在专业版权益限制；因此报告将权限阻断本身作为原型状态记录。",
        COLORS["lavender"],
    )
    doc.add_page_break()


def build_report():
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_heading("1. 采集结论", level=1)
    add_bullets(
        doc,
        [
            "ReelMate 的核心结构是“剧本生成/改编 -> 项目化制片 -> 资产沉淀/复用 -> 团队协作/积分管控”。左侧固定导航和顶部积分/会员入口贯穿所有页面。",
            "当前已登录账号可进入首页、剧本、项目、用户资产库和团队成员管理；项目内已有示例项目 try，可进入单集分镜工作台。",
            "积分与会员权益是关键产品机制：剧本生成、视频生成、团队资产库、团队席位和并发能力都与积分/专业版/扩容强绑定。",
            "复刻原型不应只做“生成器页面”，还要做权限门槛、空状态、已有项目状态、成本提示、批量操作和团队管控这些运营级状态。",
            "本版补采了页面内强相关按钮的触发态：剧本原创设定、表单校验、分集数下拉、项目创建校验、项目更多菜单/重命名、资产提取上传弹窗、视频模型选择、生成校验、添加分镜、团队成员付费门槛与成员规则说明。",
            "截图可信性：首页、剧本、项目、资产库、团队及关键弹窗均按 Chrome 窗口 ID 重新采集并目检；未实际执行消耗积分、真实上传、购买套餐或改变团队配置的动作。",
        ],
    )

    doc.add_heading("2. 信息架构与全局导航", level=1)
    add_simple_table(
        doc,
        ["主线", "可见入口/路由", "核心模块", "当前采集状态"],
        [
            ("首页", "/home", "产品入口、创建项目、积分、创作手册、商务合作", "已截图"),
            ("剧本", "/script-manage, /novel-to-script/:id", "小说改编、AI原创、剧本列表、筛选搜索、积分门槛", "已补采原创设定弹窗、填写校验和分集数下拉；深层生成受积分/实际生成限制"),
            ("项目", "/series-manage, /series-manage/:id, /series-manage/episode/:id", "项目列表、项目设置、资产准备、剧集、分镜、图片/视频生成", "已补采创建弹窗、校验、更多菜单、重命名、项目上传/导入弹窗和单集生成参数"),
            ("资产库", "/user-assets, /asset, /series-manage/asset-extract, /series-manage/asset-confirm", "历史创作、上传、提示词、团队/项目资产、资产提取确认", "用户资产库、官方/团队资产库已截图；项目内 AI提取入口已补采上传/剧本库/分镜单三类来源"),
            ("团队", "/team, /team/dashboard, /team/knowledge-base", "数据管理、成员、角色、项目权限、积分、知识库", "已补采团队首页、看板、创建成员付费门槛和成员规则说明；知识库路由当前呈现资产库视图"),
        ],
    )
    add_screenshot(
        doc,
        "00_home_logged_in.png",
        "首页 / 全局工作台入口",
        {
            "purpose": "展示已登录后的产品主入口：左侧主导航、顶部积分/会员入口、首页 hero 与“创建项目”转化入口。",
            "entry": "访问 https://www.reelmate.cn/home，账号处于登录状态。",
            "controls": "左侧导航：首页、剧本、项目、资产库、工具箱、团队；顶部 Seedance 2.0 促销、创作手册、商务合作、积分 0、购物车/客服/头像；主 CTA 创建项目。",
            "next": "用户从首页进入项目创建，或直接从左侧导航切入剧本/项目/资产库/团队。",
            "prototype": "复刻时优先保留“制片工作台”外壳：固定左侧导航、顶部积分状态、促销/帮助入口和清晰的创建项目 CTA。",
        },
    )

    doc.add_heading("3. 用户旅途总览", level=1)
    add_simple_table(
        doc,
        ["旅途", "流程串联", "结果/阻断"],
        [
            ("剧本", "首页/左侧剧本 -> 剧本管理 -> 选择 AI原创 -> 填写名称/受众/题材/集数/灵感 -> 校验通过 -> 生成规划方案", "已验证必填联动：未选拆分集数时按钮禁用；选择 40/50/60/自定义后可进入下一步"),
            ("项目", "首页/左侧项目 -> 项目列表 -> 创建/更多菜单 -> 项目详情 -> 上传剧本或分镜单 -> 剧集 -> 单集工作台", "已验证创建校验、卡片更多菜单、重命名弹窗和 AI提取资产的三类来源选择"),
            ("资产库", "左侧资产库 -> 用户资产库 -> 历史创作/上传/提示词 -> 复用至项目；项目详情 -> AI智能提取资产 -> 上传/选择剧本或分镜单 -> 确认角色/场景/道具", "用户资产库为空；项目内资产提取入口已触发到上传/剧本库/分镜单 Tab；团队资产库受专业版限制"),
            ("团队", "左侧团队 -> 团队首页 -> 数据管理/成员管理 -> 规则说明 -> 创建成员 -> 付费/席位门槛 -> 详细看板", "已验证成员创建会进入积分/专业版付费门槛；规则说明提供角色权限矩阵和积分/安全规则"),
        ],
    )

    doc.add_heading("4. 剧本旅途", level=1)
    add_note(doc, "流程图：剧本入口 -> 选择改编/原创模式 -> 输入素材或灵感 -> AI分析世界观/角色/章纲/集纲 -> 生成分集剧本 -> 进入项目或资产提取。", COLORS["gray"])
    add_screenshot(
        doc,
        "01_script_manage_empty.png",
        "剧本管理 / 创建入口与空列表",
        {
            "purpose": "承载剧本创作入口与“我的剧本”资产管理，是从文本内容进入生产链路的第一站。",
            "entry": "左侧导航点击“剧本”，进入 /script-manage。",
            "controls": "小说改编剧本：从分析开始改编小说、直接开始改编小说；AI创作剧本：从故事灵感创作剧本、从剧本创作衍生剧本；我的剧本列表、搜索、类型/排序筛选。",
            "next": "选择创建方式后进入小说/原创剧本流程；已有剧本可进入编辑、查看或后续项目/资产链路。",
            "prototype": "卡片区要体现两类创作路径和能力差异；列表区必须支持空状态、搜索筛选、排序、剧本数量统计和未来的批量/状态操作。",
        },
    )
    add_screenshot(
        doc,
        "12_script_ai_original_settings_modal.png",
        "剧本创建 / AI原创剧本设定弹窗",
        {
            "purpose": "承接“从故事灵感创作剧本”的参数设定，是 AI原创剧本从灵感进入规划方案的第一步。",
            "entry": "剧本管理 -> AI创作剧本 -> 从故事灵感创作剧本。",
            "controls": "文件名称、剧本受众、题材看点、拆分集数、分卡设置、每集长度、创作灵感输入框、积分详情、取消、完成设定生成规划方案；必填项未完成时生成按钮禁用。",
            "next": "补齐剧本名称、分集数和灵感文本后生成规划方案，再进入世界观/角色/章纲/集纲/分集剧本链路。",
            "prototype": "创建设定弹窗要表现表单依赖、自动适配选项、积分成本、禁用态和限时特惠标签，这比单一 prompt 输入更接近真实产品。",
        },
    )
    add_screenshot(
        doc,
        "13_script_ai_original_filled_disabled.png",
        "剧本创建 / 填写后仍需拆分集数的禁用态",
        {
            "purpose": "验证原创剧本不是单字段 prompt，而是多字段表单联动；名称和灵感已填后，缺少拆分集数仍不能提交。",
            "entry": "在 AI原创剧本设定弹窗中填写文件名称和创作灵感，保留拆分集数为空。",
            "controls": "文件名称字符计数、剧本受众、题材看点、拆分集数必填提示、创作灵感输入框、积分详情 14、禁用的“完成设定，生成规划方案”。",
            "next": "选择拆分集数后生成按钮才可进入规划方案生成。",
            "prototype": "必须复刻表单校验与按钮禁用逻辑；错误提示应靠近触发字段，底部积分成本随关键参数保持可见。",
        },
    )
    add_screenshot(
        doc,
        "14_script_episode_count_dropdown.png",
        "剧本创建 / 拆分集数下拉选项",
        {
            "purpose": "展示剧本拆分规模的核心参数，直接影响后续规划、成本和生成量。",
            "entry": "点击 AI原创剧本设定弹窗中的“拆分集数”。",
            "controls": "下拉项包括 40集、50集、60集、自定义分集（1-100）；选择后会回填字段并解锁下一步。",
            "next": "选择集数 -> 生成规划方案 -> 进入世界观、角色、章纲、集纲和分集剧本。",
            "prototype": "分集数应作为结构化字段存储，不应只写进 prompt；它会驱动剧集数量、成本估算、进度和项目初始化。",
        },
    )
    add_bullets(
        doc,
        [
            "公开前端路由显示深层剧本流程包含小说内容分析、世界观编辑、角色卡编辑、章纲情节编辑、集纲/分集生成。",
            "剧本章节模型建议拆为 Script、WorldSetting、Character、ChapterOutline、EpisodeOutline、EpisodeScript，便于后续资产提取和项目引用。",
        ],
    )

    doc.add_heading("5. 项目旅途", level=1)
    add_note(doc, "流程图：项目列表 -> 项目详情总览 -> 项目资产/剧集 -> 单集分镜工作台 -> 分镜图生成 -> 视频生成/批量操作 -> 导出或同步后期。", COLORS["gray"])
    add_screenshot(
        doc,
        "03_project_list_with_existing_project.png",
        "项目列表 / 项目管理入口",
        {
            "purpose": "管理所有漫剧项目，并作为进入制片链路的入口。",
            "entry": "左侧导航点击“项目”，进入 /series-manage。",
            "controls": "全部项目数量、状态筛选、搜索、项目卡片、创建项目按钮；项目卡显示名称、创建时间、项目状态。",
            "next": "点击项目卡进入项目详情；点击创建项目进入新项目配置。",
            "prototype": "项目卡应展示最少信息集：封面/名称/状态/时间/进度/成员或成本摘要；列表要支持搜索、状态筛选和空/有数据两种状态。",
        },
    )
    add_screenshot(
        doc,
        "15_project_create_modal.png",
        "项目列表 / 创建项目弹窗",
        {
            "purpose": "验证“创建项目”不是直接跳转，而是先配置项目基础属性。",
            "entry": "项目列表 -> 点击右上角“创建项目”。",
            "controls": "项目名称、画面比例 9:16/16:9、剧目类型卡片：国内仿真人剧、海外仿真人剧、2D/3D动漫、确认按钮。",
            "next": "填写项目名称并选择比例/类型后创建项目，进入项目详情或项目列表。",
            "prototype": "项目创建应优先收集影响后续生成规格的字段：名称、比例、剧目类型；类型卡片要用视觉差异区分真实人剧和动漫项目。",
        },
    )
    add_screenshot(
        doc,
        "16_project_create_validation_toast.png",
        "项目列表 / 创建项目必填校验",
        {
            "purpose": "补足创建动作的失败反馈，说明空表单提交不会静默失败。",
            "entry": "在创建项目弹窗不填写项目名称/画面比例时点击“确认”。",
            "controls": "页面顶部 toast：请填写项目名称和画面比例；弹窗保持打开，用户可继续补填。",
            "next": "补齐必填项后再次确认。",
            "prototype": "复刻时要有统一 toast/消息系统；创建类弹窗需保留用户上下文，不要因校验失败关闭。",
        },
    )
    add_screenshot(
        doc,
        "17_project_card_more_menu.png",
        "项目列表 / 项目卡更多菜单",
        {
            "purpose": "展示项目卡片上的轻量管理能力，而不是只有进入详情一种动作。",
            "entry": "在项目 try 卡片点击右上角更多按钮。",
            "controls": "菜单项：上传封面、重命名、删除。",
            "next": "选择上传封面进入素材上传；选择重命名打开名称编辑；选择删除进入确认风险弹窗。",
            "prototype": "项目卡要提供行内操作菜单；危险操作和非危险操作需分层展示，删除必须二次确认。",
        },
    )
    add_screenshot(
        doc,
        "18_project_rename_modal.png",
        "项目列表 / 项目重命名弹窗",
        {
            "purpose": "验证项目元数据可被快速编辑。",
            "entry": "项目卡更多菜单 -> 重命名。",
            "controls": "新项目名称输入框、当前值 try、字符计数 3/50、取消、保存。",
            "next": "保存后回到项目列表并刷新卡片名称。",
            "prototype": "名称编辑要支持原值回填、字符限制、保存/取消；建议复刻为通用的“单字段编辑弹窗”组件。",
        },
    )
    add_screenshot(
        doc,
        "05_project_overview_assets_and_episodes.png",
        "项目详情 / 总览、资产准备与剧集",
        {
            "purpose": "连接项目配置、资产准备、AI资产提取和剧集生产，是项目制片的中心页。",
            "entry": "项目列表 -> 项目 try -> 总览。",
            "controls": "项目名称 try、状态未开始、2D/3D动漫、9:16、风格信息；上传剧本/分镜入口；角色/场景/道具/其他资产卡；AI智能提取资产；剧集卡 Try/test。",
            "next": "进入 AI智能提取资产确认清单，或打开单集进入分镜工作台。",
            "prototype": "总览页要把“项目设置、资产准备、剧集生产”放在一屏串联；资产卡应显示数量/状态，剧集卡应支持进入、状态、成本/进度。",
        },
    )
    add_screenshot(
        doc,
        "19_project_script_upload_modal_from_extract.png",
        "项目详情 / AI智能提取资产的剧本上传入口",
        {
            "purpose": "点击“AI智能提取资产”后，系统先要求确定素材来源，说明资产提取是项目内的导入/解析流程。",
            "entry": "项目详情 -> 资产准备 -> AI智能提取资产。",
            "controls": "弹窗 Tab：剧本库、剧本上传、分镜单上传；剧本上传区支持点击或拖拽 docx/txt；确认上传按钮。",
            "next": "上传剧本后进入角色、场景、道具等资产提取清单和确认页。",
            "prototype": "资产提取入口应被设计成向导：选择来源 -> 上传/选择素材 -> 解析 -> 确认资产 -> 回流项目资产库。",
        },
    )
    add_screenshot(
        doc,
        "20_project_script_library_empty_in_upload_modal.png",
        "项目详情 / AI提取资产的剧本库空状态",
        {
            "purpose": "展示不上传文件时可从已有剧本库选择，但当前账号没有可用剧本。",
            "entry": "AI智能提取资产弹窗 -> 切换到“剧本库”。",
            "controls": "搜索输入框、空状态插画/提示、确认按钮不可推进。",
            "next": "选择已有剧本后进入资产提取；若无剧本则回到上传或先去剧本模块创建。",
            "prototype": "跨模块复用要有清晰空状态和回流入口；剧本库选择器应支持搜索、列表、选择态和空态引导。",
        },
    )
    add_screenshot(
        doc,
        "21_project_storyboard_upload_tab.png",
        "项目详情 / 分镜单上传入口",
        {
            "purpose": "说明项目可以不从剧本文本开始，也能从已有分镜单导入进入生产。",
            "entry": "AI智能提取资产弹窗 -> 切换到“分镜单上传”。",
            "controls": "文本样式/表格样式分镜单说明、下载模板、点击放大查看、上传区支持 doc/docx/txt/xls/xlsx、确认上传。",
            "next": "上传分镜单后进入分镜/资产解析和项目剧集生产。",
            "prototype": "导入能力要支持模板下载和格式说明；这类入口能显著降低从外部制片文档迁移到平台的成本。",
        },
    )
    add_screenshot(
        doc,
        "06_episode_storyboard_editor_guided_tour.png",
        "单集工作台 / 分镜与视频生成",
        {
            "purpose": "单集生产的核心操作台，用于 AI拆分镜、分镜图管理和视频生成。",
            "entry": "项目详情 -> 剧集卡 -> /series-manage/episode/:id。",
            "controls": "AI拆分镜引导、分镜数、添加分镜、进入时间轴；分镜图片/分镜视频 Tab；右侧模型选择 Vidu Q3-Pro、首帧/首尾帧/参考图/AI改视频、提示词、音画同步、时长、分辨率、消耗积分和生成按钮。",
            "next": "用户生成/编辑分镜图片，再用分镜图作为首帧或参考生成视频；可进入时间轴做后续编排。",
            "prototype": "这是复刻优先级最高的工作台：左侧分镜列表、中间预览/分镜区、右侧生成参数面板、积分成本、引导层和批量/时间轴入口都应实现。",
        },
    )
    add_screenshot(
        doc,
        "22_episode_video_model_dropdown.png",
        "单集工作台 / 视频模型下拉选择",
        {
            "purpose": "展示视频生成不是单一模型，而是带能力标签和价格差异的模型市场式选择。",
            "entry": "单集工作台 -> 右侧分镜视频参数 -> 点击模型下拉。",
            "controls": "Happy Horse、Vidu Q3-Pro、Vidu Q2、即梦 3.0/3.5 Pro、Hailuo 2.3 等模型；能力标签包括时长、首尾帧、生成音频、口型同步等；底部显示积分消耗与生成按钮。",
            "next": "选择模型后配置首帧/尾帧/参考图/提示词/时长/分辨率并提交生成。",
            "prototype": "模型选择器要复刻能力标签、默认选中态、成本联动和模型兼容性限制；这是用户决定成本和效果的关键控件。",
        },
    )
    add_screenshot(
        doc,
        "23_episode_generate_validation_toast.png",
        "单集工作台 / 立即生成的首帧校验",
        {
            "purpose": "点击“立即生成”后触发表单校验，说明生成任务必须满足素材条件。",
            "entry": "单集工作台 -> 未上传首帧图时点击“立即生成”。",
            "controls": "toast/提示：请上传完毕首帧图后提交生成任务；右侧生成按钮附近保留参数上下文。",
            "next": "上传首帧图或切换生成模式后重新提交。",
            "prototype": "生成工作台必须有任务提交前校验；错误提示要指向缺失素材，避免用户消耗积分前产生失败任务。",
        },
    )
    add_screenshot(
        doc,
        "24_episode_add_storyboard_result.png",
        "单集工作台 / 添加分镜后的列表状态",
        {
            "purpose": "验证“添加分镜”会改变分镜列表结构，而不只是静态按钮。",
            "entry": "单集工作台 -> 点击“添加分镜”。",
            "controls": "分镜数量从 2 变为 3；新增分镜卡片 3，状态为未定稿；中间预览区和右侧参数面板保持可继续编辑。",
            "next": "为新分镜补图、补提示词或进入视频生成。",
            "prototype": "分镜列表应支持增删改、排序、状态标记和选中态；新增分镜后要保持用户在当前工作台内连续创作。",
        },
    )

    doc.add_heading("6. 资产库旅途", level=1)
    add_note(doc, "流程图：资产库 -> 历史创作/上传/提示词 -> 搜索筛选/收藏/批量 -> 复用到项目；项目详情 -> AI智能提取资产 -> 角色/场景/道具确认 -> 回流项目资产。", COLORS["gray"])
    add_screenshot(
        doc,
        "07_user_asset_library_empty.png",
        "用户资产库 / 历史创作与上传资产",
        {
            "purpose": "沉淀个人历史生成、上传内容和提示词，是跨项目复用素材的入口。",
            "entry": "左侧导航点击“资产库”，当前进入 /user-assets。",
            "controls": "历史创作、Agent项目、历史上传、我的提示词 Tab；类型筛选、搜索、我的收藏、批量操作、时间顺序、文件夹；空状态提示。",
            "next": "有资产后可筛选、收藏、批量管理或复用到项目；上传/生成资产会回流到此处。",
            "prototype": "资产库必须支持空状态和列表状态；数据维度至少包括类型、来源、项目归属、收藏、文件夹、创建时间、生成参数。",
        },
    )
    add_screenshot(
        doc,
        "10_team_knowledge_base_gate.png",
        "官方/团队资产库 / 角色、场景、道具分类",
        {
            "purpose": "展示平台预置资产与团队复用资产的浏览入口，当前路由 /team/knowledge-base 实际呈现为资产库视图。",
            "entry": "直接访问 /team/knowledge-base，页面落在左侧“资产库”主导航。",
            "controls": "官方资产库与团队资产库 Tab、角色/场景/道具分类、左侧文件夹树、资产卡片网格、搜索框、顶部积分与会员入口。",
            "next": "点击资产卡进入详情/复用，切换团队资产库查看团队沉淀内容；专业版团队资产库需权限开通。",
            "prototype": "资产库原型要同时覆盖“个人历史资产”和“官方/团队素材库”：前者偏管理，后者偏浏览、分类、复用和权限门槛。",
        },
    )
    add_bullets(
        doc,
        [
            "项目总览截图中已有项目级资产准备：角色、场景、道具、其他，以及“AI智能提取资产（首次免费）”。这说明资产库不是孤立模块，而是和剧本/项目联动。",
            "公开路由包含 /series-manage/asset-extract 与 /series-manage/asset-confirm；直接访问 asset-extract 当前只显示空壳和返回按钮，推断需要项目/剧本上下文才能渲染提取清单。",
            "团队资产库入口在团队页出现专业版权益提示，当前账号无法直接使用；复刻时需要实现“团队资产库未开通”与“个人资产库可用”的权限差异。",
        ],
    )

    doc.add_heading("7. 团队旅途", level=1)
    add_note(doc, "流程图：团队首页 -> 数据管理/详细看板 -> 成员管理 -> 创建成员/角色/项目/成员组/状态 -> 积分分配与记录 -> 团队知识库。", COLORS["gray"])
    add_screenshot(
        doc,
        "08_team_member_management_gate.png",
        "团队 / 权益门槛、数据管理与成员管理",
        {
            "purpose": "团队协作与运营管控入口，承载专业版权益、团队数据、成员账号和积分管理。",
            "entry": "左侧导航点击“团队”，进入 /team。",
            "controls": "专业版团队资产库提示与去开通按钮；数据管理卡：团队项目、团队席位、单账号并发、消耗/剩余/可分配积分；查看详细数据看板、刷新；成员管理搜索项：账号、成员名称、角色、项目、状态、备注；创建成员账号。",
            "next": "开通/扩容/加量后启用团队资产库、席位与积分；创建成员后进入成员表操作；点击详细数据看板进入 /team/dashboard。",
            "prototype": "团队页要有管理员视角的空状态和权限提示；成员表字段、筛选条件、积分/席位/并发指标都应作为原型必备。",
        },
    )
    add_screenshot(
        doc,
        "26_team_member_rules_modal.png",
        "团队 / 成员管理规则说明",
        {
            "purpose": "补足团队成员体系的规则层：角色、权限、成员组、积分和账号安全都在此弹窗中解释。",
            "entry": "团队首页 -> 成员管理标题旁点击“规则说明”。",
            "controls": "关闭、确认；基础规则、成员角色权限管理、角色权限对照表、成员组管理、积分管理机制、账号与安全管理。",
            "next": "理解权限后创建成员账号、分配角色/项目/成员组/积分。",
            "prototype": "团队原型需要把角色权限矩阵数据化：管理员、组管理员、导演、动画师、编剧、剪辑师等角色对应不同创作、下载、删除、知识库和团队管理权限。",
        },
    )
    add_screenshot(
        doc,
        "25_team_create_member_pricing_gate.png",
        "团队 / 创建成员触发套餐与积分门槛",
        {
            "purpose": "验证创建成员并不总是进入成员表单；当前账号因席位/会员权益受限先进入购买或加量弹窗。",
            "entry": "团队首页 -> 点击“创建成员账号”。",
            "controls": "积分加量/兑换码 Tab；体验版、专业版、企业版套餐卡；购买/联系商务入口；弹窗关闭按钮。",
            "next": "购买专业版/扩容席位后才可继续创建子账号和分配权限。",
            "prototype": "成员创建要实现权限和商业化前置校验；受限态应明确说明原因，并提供开通、扩容、兑换码或商务联系路径。",
        },
    )
    add_screenshot(
        doc,
        "09_team_dashboard_gate.png",
        "团队详细数据看板 / 成员创作与消耗",
        {
            "purpose": "提供团队维度的成员创作量、项目资产成本和排行榜分析，是管理员做积分和产能管理的页面。",
            "entry": "团队首页 -> 查看详细数据看板，或直接访问 /team/dashboard。",
            "controls": "顶部返回、成员创作与消耗/项目资产与成本/排行榜 Tab；成员数、启用成员数、成员总消耗积分、成员均消耗积分；角色/状态筛选、日期范围、今天/昨天/本周/本月/上月/今年快捷筛选、成员积分管理明细、导出。",
            "next": "筛选时间或角色查看成员消耗明细，导出数据，或进入成员积分管理明细做分配/追踪。",
            "prototype": "团队看板要服务运营管理：指标卡、筛选条、明细表、导出、空状态都要齐；数据模型要能按成员、角色、项目、时间和积分消耗聚合。",
        },
    )
    add_bullets(
        doc,
        [
            "前端权限注释显示角色层级包括主账号、管理员、组管理员、导演、动画师、编剧；需要在原型中体现入口差异、禁用按钮和无权限提示。",
            "团队知识库路由 /team/knowledge-base 当前呈现为官方/团队资产库视图，说明平台把团队知识沉淀和资产复用放在同一类导航心智下；真实知识文档能力仍需二次确认。",
        ],
    )

    doc.add_heading("8. 原型需求清单", level=1)
    add_simple_table(
        doc,
        ["优先级", "页面/状态", "复刻要求"],
        [
            ("必做", "全局框架", "左侧导航、顶部积分/会员/帮助入口、用户头像、深色主题与紫色高亮。"),
            ("必做", "剧本管理", "创建模式卡片、我的剧本列表、空状态、筛选搜索、原创设定弹窗、必填校验、分集数下拉、积分成本。"),
            ("必做", "项目列表与详情", "项目卡、状态筛选、创建项目弹窗、创建校验、卡片更多菜单、重命名、详情二级导航、资产准备、剧集卡。"),
            ("必做", "单集工作台", "分镜列表、添加分镜、AI拆分镜、图片/视频 Tab、模型选择器、提示词/参数/积分成本/生成校验/引导层。"),
            ("必做", "资产库", "个人资产库 Tab、搜索筛选、收藏、批量、文件夹、空状态、项目资产分类、AI提取资产来源选择、剧本/分镜单上传。"),
            ("必做", "团队", "专业版权益门槛、数据管理指标、成员表、成员规则说明、成员创建付费门槛、角色/项目/状态筛选。"),
            ("可延后", "支付与兑换", "套餐购买、兑换码、发票/订单链路可先做高保真弹窗，不接真实支付。"),
            ("可延后", "详细看板/知识库", "先做信息架构和入口，后续补充真实团队数据、知识文档和权限协作。"),
        ],
    )
    add_simple_table(
        doc,
        ["数据模型", "关键字段"],
        [
            ("User / Account", "id, name, avatar, role, creditBalance, membershipTier"),
            ("Team / Member", "teamId, seats, concurrency, members, memberGroup, role, status, assignedProjects"),
            ("Project", "name, status, style, aspectRatio, animationType, owner, members, assetCounts, episodeCounts"),
            ("Script", "type, sourceText, analysisStatus, worldSetting, characters, chapterOutlines, episodeOutlines"),
            ("Episode / Storyboard", "episodeId, storyboardIndex, prompt, image, video, model, duration, resolution, generationStatus"),
            ("Asset", "type(role/scene/prop/other), source, projectId, folder, favorite, generationParams, reusableScope"),
            ("CreditTransaction", "scene, amount, beforeBalance, afterBalance, operator, createdAt"),
        ],
    )

    doc.add_heading("9. 附录：路由与权限观察", level=1)
    add_simple_table(
        doc,
        ["路由", "模块", "说明"],
        [
            ("/home", "首页", "已登录后仍以创建项目作为主 CTA。"),
            ("/script-manage", "剧本", "剧本创建入口与我的剧本列表。"),
            ("/novel-to-script/:id", "剧本", "小说改编深层流程，预计包含分析、世界观、角色、章纲、集纲、分集。"),
            ("/series-manage", "项目", "项目列表与创建入口。"),
            ("/series-manage/:id", "项目", "项目详情总览、资产、剧集、成员、统计。"),
            ("/series-manage/episode/:id", "项目", "单集分镜、图片与视频生成工作台。"),
            ("/series-manage/asset-extract", "资产提取", "直接访问显示空壳和返回按钮，需项目上下文二次确认。"),
            ("/series-manage/asset-confirm", "资产确认", "角色/场景/道具确认与批量生成页，需上下文二次确认。"),
            ("/user-assets", "资产库", "个人历史创作、上传和提示词。"),
            ("/asset", "团队资产库", "专业版会员权益门槛。"),
            ("/team", "团队", "团队数据、成员管理、积分/席位/并发入口。"),
            ("/team/dashboard", "团队", "详细数据看板入口，可从团队页点击进入。"),
            ("/team/knowledge-base", "团队/资产", "当前渲染为官方/团队资产库视图，真实知识文档能力需二次确认。"),
        ],
    )

    doc.add_heading("10. 未知项与二次确认", level=1)
    add_bullets(
        doc,
        [
            "当前账号积分为 0，因此未实际执行会消耗积分的剧本生成、图片生成、视频生成；本版已补采生成前的参数、下拉、校验和成本状态。",
            "未真实上传剧本/分镜单、购买套餐或创建成员账号，避免改动线上账号数据；上传后的解析结果、资产确认清单和成员表新增状态建议用专门测试账号二次补采。",
            "团队资产库明确为专业版会员权益；当前团队席位、并发、剩余积分均为 0，已验证创建成员会进入套餐/积分门槛，但无法验证完整成员创建后的表格状态和积分分配记录。",
            "资产提取/确认深层页需要有效项目/剧本文本上下文；本版已通过项目总览入口补采剧本库、剧本上传、分镜单上传三类来源，解析后的角色/场景/道具确认页仍需二次确认。",
            "报告未修改线上真实项目数据；采集使用已有项目 try 与现有空状态，避免误消耗积分或改变团队配置。",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End of report")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
